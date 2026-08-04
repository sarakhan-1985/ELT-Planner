
import streamlit as st

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle
)

from prompts import build_prompt
from lesson_generator import generate_lesson


# --------------------------------------------------
# HELPER FUNCTIONS
# --------------------------------------------------

def clean_markdown_text(text):
    """
    Remove common Markdown formatting for Word and PDF files.
    """
    if not text:
        return ""

    replacements = {
        "**": "",
        "__": "",
        "`": "",
        "### ": "",
        "## ": "",
        "# ": ""
    }

    cleaned_text = text

    for old, new in replacements.items():
        cleaned_text = cleaned_text.replace(old, new)

    return cleaned_text.strip()


def is_markdown_separator_row(cells):
    """
    Check whether a Markdown table row is a separator row.
    Example:
    |---|---|---|
    """
    return all(
        cell.replace("-", "").replace(":", "").strip() == ""
        for cell in cells
    )


def markdown_to_word(lesson_text):
    """
    Convert generated Markdown lesson-plan content into a Word document.
    """
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title = document.add_heading("ELT Lesson Plan", level=0)
    title.alignment = 1

    lines = lesson_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].strip()

        if not line:
            index += 1
            continue

        # ------------------------------------------
        # MARKDOWN TABLE
        # ------------------------------------------

        if (
            line.startswith("|")
            and line.endswith("|")
            and line.count("|") >= 2
        ):
            table_lines = []

            while index < len(lines):
                current_line = lines[index].strip()

                if (
                    current_line.startswith("|")
                    and current_line.endswith("|")
                ):
                    table_lines.append(current_line)
                    index += 1
                else:
                    break

            rows = []

            for table_line in table_lines:
                cells = [
                    clean_markdown_text(cell.strip())
                    for cell in table_line.strip("|").split("|")
                ]

                if not is_markdown_separator_row(cells):
                    rows.append(cells)

            if rows:
                column_count = max(len(row) for row in rows)

                word_table = document.add_table(
                    rows=1,
                    cols=column_count
                )

                word_table.style = "Table Grid"

                # Header row
                header_cells = word_table.rows[0].cells

                for column_index in range(column_count):
                    header_text = (
                        rows[0][column_index]
                        if column_index < len(rows[0])
                        else ""
                    )

                    header_cells[column_index].text = header_text

                    for paragraph in header_cells[column_index].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)

                # Data rows
                for row_data in rows[1:]:
                    row_cells = word_table.add_row().cells

                    for column_index in range(column_count):
                        cell_text = (
                            row_data[column_index]
                            if column_index < len(row_data)
                            else ""
                        )

                        row_cells[column_index].text = cell_text

                        for paragraph in row_cells[column_index].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

                document.add_paragraph()

            continue

        # ------------------------------------------
        # HEADINGS
        # ------------------------------------------

        if line.startswith("### "):
            document.add_heading(
                clean_markdown_text(line),
                level=3
            )

        elif line.startswith("## "):
            document.add_heading(
                clean_markdown_text(line),
                level=2
            )

        elif line.startswith("# "):
            heading_text = clean_markdown_text(line)

            if heading_text.lower() != "lesson plan":
                document.add_heading(
                    heading_text,
                    level=1
                )

        # ------------------------------------------
        # BULLET POINTS
        # ------------------------------------------

        elif line.startswith("- ") or line.startswith("* "):
            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.add_run(
                clean_markdown_text(line[2:])
            )

        # ------------------------------------------
        # NUMBERED LISTS
        # ------------------------------------------

        elif (
            len(line) > 2
            and line[0].isdigit()
            and "." in line[:4]
        ):
            paragraph = document.add_paragraph(
                style="List Number"
            )

            content = line.split(".", 1)[1].strip()

            paragraph.add_run(
                clean_markdown_text(content)
            )

        # ------------------------------------------
        # NORMAL TEXT
        # ------------------------------------------

        else:
            paragraph = document.add_paragraph()

            paragraph.add_run(
                clean_markdown_text(line)
            )

        index += 1

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output.getvalue()


def markdown_to_pdf(lesson_text):
    """
    Convert generated Markdown lesson-plan content into a PDF.
    """
    output = BytesIO()

    pdf_document = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=24,
        leftMargin=24,
        topMargin=24,
        bottomMargin=24
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "LessonTitle",
        parent=styles["Title"],
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=12
    )

    heading_style = ParagraphStyle(
        "LessonHeading",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=8,
        spaceAfter=5
    )

    body_style = ParagraphStyle(
        "LessonBody",
        parent=styles["BodyText"],
        fontSize=8,
        leading=11,
        spaceAfter=4
    )

    bullet_style = ParagraphStyle(
        "LessonBullet",
        parent=body_style,
        leftIndent=14,
        firstLineIndent=-8,
        bulletIndent=5
    )

    table_header_style = ParagraphStyle(
        "TableHeader",
        parent=body_style,
        fontSize=7,
        leading=9,
        alignment=TA_CENTER,
        textColor=colors.white
    )

    table_cell_style = ParagraphStyle(
        "TableCell",
        parent=body_style,
        fontSize=6.5,
        leading=8
    )

    story = [
        Paragraph("ELT Lesson Plan", title_style),
        Spacer(1, 8)
    ]

    lines = lesson_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].strip()

        if not line:
            index += 1
            continue

        # ------------------------------------------
        # MARKDOWN TABLE
        # ------------------------------------------

        if (
            line.startswith("|")
            and line.endswith("|")
            and line.count("|") >= 2
        ):
            table_lines = []

            while index < len(lines):
                current_line = lines[index].strip()

                if (
                    current_line.startswith("|")
                    and current_line.endswith("|")
                ):
                    table_lines.append(current_line)
                    index += 1
                else:
                    break

            rows = []

            for table_line in table_lines:
                cells = [
                    clean_markdown_text(cell.strip())
                    for cell in table_line.strip("|").split("|")
                ]

                if not is_markdown_separator_row(cells):
                    rows.append(cells)

            if rows:
                column_count = max(len(row) for row in rows)

                formatted_rows = []

                for row_index, row in enumerate(rows):
                    paragraph_style = (
                        table_header_style
                        if row_index == 0
                        else table_cell_style
                    )

                    padded_row = row + [""] * (
                        column_count - len(row)
                    )

                    formatted_rows.append([
                        Paragraph(cell, paragraph_style)
                        for cell in padded_row
                    ])

                available_width = landscape(A4)[0] - 48

                if column_count == 5:
                    column_widths = [
                        available_width * 0.08,
                        available_width * 0.30,
                        available_width * 0.30,
                        available_width * 0.20,
                        available_width * 0.12
                    ]

                elif column_count == 4:
                    column_widths = [
                        available_width * 0.18,
                        available_width * 0.34,
                        available_width * 0.28,
                        available_width * 0.20
                    ]

                else:
                    column_widths = [
                        available_width / column_count
                    ] * column_count

                pdf_table = Table(
                    formatted_rows,
                    colWidths=column_widths,
                    repeatRows=1
                )

                pdf_table.setStyle(
                    TableStyle([
                        (
                            "BACKGROUND",
                            (0, 0),
                            (-1, 0),
                            colors.HexColor("#003366")
                        ),
                        (
                            "TEXTCOLOR",
                            (0, 0),
                            (-1, 0),
                            colors.white
                        ),
                        (
                            "FONTNAME",
                            (0, 0),
                            (-1, 0),
                            "Helvetica-Bold"
                        ),
                        (
                            "VALIGN",
                            (0, 0),
                            (-1, -1),
                            "TOP"
                        ),
                        (
                            "GRID",
                            (0, 0),
                            (-1, -1),
                            0.5,
                            colors.grey
                        ),
                        (
                            "ROWBACKGROUNDS",
                            (0, 1),
                            (-1, -1),
                            [
                                colors.white,
                                colors.HexColor("#F2F5F8")
                            ]
                        ),
                        (
                            "LEFTPADDING",
                            (0, 0),
                            (-1, -1),
                            4
                        ),
                        (
                            "RIGHTPADDING",
                            (0, 0),
                            (-1, -1),
                            4
                        ),
                        (
                            "TOPPADDING",
                            (0, 0),
                            (-1, -1),
                            4
                        ),
                        (
                            "BOTTOMPADDING",
                            (0, 0),
                            (-1, -1),
                            4
                        )
                    ])
                )

                story.append(pdf_table)
                story.append(Spacer(1, 10))

            continue

        # ------------------------------------------
        # HEADINGS
        # ------------------------------------------

        if line.startswith("#"):
            heading_text = clean_markdown_text(line)

            if heading_text.lower() != "lesson plan":
                story.append(
                    Paragraph(
                        heading_text,
                        heading_style
                    )
                )

        # ------------------------------------------
        # BULLET POINTS
        # ------------------------------------------

        elif line.startswith("- ") or line.startswith("* "):
            story.append(
                Paragraph(
                    "• " + clean_markdown_text(line[2:]),
                    bullet_style
                )
            )

        # ------------------------------------------
        # NORMAL TEXT
        # ------------------------------------------

        else:
            story.append(
                Paragraph(
                    clean_markdown_text(line),
                    body_style
                )
            )

        index += 1

    pdf_document.build(story)

    output.seek(0)

    return output.getvalue()


# --------------------------------------------------
# PAGE CONFIGURATION
# --------------------------------------------------

st.set_page_config(
    page_title="ELT Planner AI",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)


# --------------------------------------------------
# CUSTOM CSS
# --------------------------------------------------

st.markdown(
    """
    <style>

    .main {
        background-color: #F8FAFC;
    }

    h1 {
        color: #003366;
        font-weight: 700;
    }

    h2 {
        color: #004B8D;
    }

    .section-box {
        background-color: white;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0px 2px 10px rgba(0,0,0,0.08);
        margin-bottom: 20px;
    }

    .stButton > button {
        background-color: #003366;
        color: white;
        border-radius: 10px;
        min-height: 55px;
        font-size: 18px;
        font-weight: bold;
    }

    .stButton > button:hover {
        background-color: #00509E;
        color: white;
    }

    div[data-testid="stDownloadButton"] button {
        min-height: 50px;
        font-size: 16px;
        font-weight: 600;
    }

    </style>
    """,
    unsafe_allow_html=True
)


# --------------------------------------------------
# HOME PAGE HEADER
# --------------------------------------------------

st.title("🎓 ELT Planner AI")

st.markdown(
    """
    ### Professional Lesson Planning Assistant for English Language Teachers

    Design engaging, Bloom's Taxonomy-based lesson plans in minutes.

    Upload your Course Outline or enter your Course Objectives and CLOs,
    select your programme, choose your lesson topic, and let AI generate
    a professional lesson plan.
    """
)

st.divider()


# --------------------------------------------------
# STEP 1 — COURSE INFORMATION
# --------------------------------------------------

st.markdown("## 📘 Step 1 — Course Information")

col1, col2 = st.columns(2)

with col1:
    course = st.selectbox(
        "Course Title",
        [
            "Functional English",
            "Expository Writing",
            "English I",
            "Technical Business Writing",
            "Communication Skills",
            "Presentation Skills"
        ]
    )

    programme = st.selectbox(
        "Programme",
        [
            "Computer Science",
            "Electrical Engineering",
            "Civil Engineering",
            "Business Administration",
            "Accounting & Finance"
        ]
    )

    year = st.selectbox(
        "Undergraduate Level",
        [
            "First Year",
            "Second Year",
            "Third Year",
            "Fourth Year",
            "Fifth Year"
        ]
    )

with col2:
    duration = st.selectbox(
        "Lesson Duration",
        [
            "40 Minutes",
            "60 Minutes",
            "90 Minutes",
            "110 Minutes"
        ]
    )

    class_size = st.number_input(
        "Class Size",
        min_value=5,
        max_value=150,
        value=30
    )

    proficiency = st.selectbox(
        "Student English Proficiency",
        [
            "Beginner",
            "Intermediate",
            "Advanced",
            "Mixed Ability"
        ]
    )

st.divider()


# --------------------------------------------------
# STEP 2 — CURRICULUM ALIGNMENT
# --------------------------------------------------

st.markdown("## 📄 Step 2 — Curriculum Alignment")

uploaded_file = st.file_uploader(
    "Upload Course Outline (PDF or DOCX)",
    type=["pdf", "docx"]
)

if uploaded_file is not None:
    st.success(
        f"File uploaded successfully: {uploaded_file.name}"
    )

st.markdown("### OR")

course_objectives = st.text_area(
    "Paste Course Objectives",
    height=180,
    placeholder="""
Example

CO1:
Develop academic reading skills.

CO2:
Develop effective academic writing.

CO3:
Develop communication skills.
"""
)

clos = st.text_area(
    "Paste Course Learning Outcomes (CLOs)",
    height=180,
    placeholder="""
Example

CLO1:
Identify important features of academic texts.

CLO2:
Apply appropriate language strategies.

CLO3:
Evaluate written communication for clarity and accuracy.
"""
)

st.divider()


# --------------------------------------------------
# STEP 3 — LESSON DETAILS
# --------------------------------------------------

st.markdown("## 📝 Step 3 — Lesson Details")

topics = {
    "Writing": [
        "Paraphrasing",
        "Summary Writing",
        "Expository Writing",
        "Narrative Writing",
        "Descriptive Writing",
        "Compare and Contrast Essay",
        "Cause and Effect Essay",
        "Persuasive Essay",
        "Email Writing",
        "Report Writing",
        "Technical Writing"
    ],

    "Reading": [
        "Stated Main Idea",
        "Implied Main Idea",
        "Supporting Details",
        "Critical Response Questions",
        "Inference",
        "Fact vs Opinion",
        "Author's Purpose",
        "Tone",
        "Vocabulary in Context"
    ],

    "Listening": [
        "Listening for Gist",
        "Listening for Details",
        "Lecture Comprehension",
        "Note Taking",
        "TED Talk Analysis"
    ],

    "Speaking": [
        "Presentation Skills",
        "Group Discussion",
        "Debate",
        "Interview Skills",
        "Extempore Speaking"
    ]
}

col1, col2 = st.columns(2)

with col1:
    skill = st.selectbox(
        "Language Skill",
        list(topics.keys())
    )

with col2:
    topic = st.selectbox(
        "Lesson Topic",
        topics[skill]
    )


# --------------------------------------------------
# ACTIVITIES TO GENERATE
# --------------------------------------------------

st.markdown("### 🎯 Activities to Generate")

activity_col1, activity_col2 = st.columns(2)

with activity_col1:
    beginner_activity = st.checkbox(
        "🟢 Beginner Activity",
        value=True
    )

    intermediate_activity = st.checkbox(
        "🟡 Intermediate Activity",
        value=True
    )

    advanced_activity = st.checkbox(
        "🔴 Advanced Activity",
        value=True
    )

with activity_col2:
    activity_bullet_points = st.checkbox(
        "• Present differentiated activities in bullet points",
        value=True,
        help=(
            "Each differentiated activity will be shown using "
            "clear but sufficiently detailed bullet points."
        )
    )

    detailed_table = st.checkbox(
        "📋 Keep lesson procedure detailed in table form",
        value=True,
        help=(
            "Teacher actions, student actions, resources and "
            "assessment instructions will remain detailed."
        )
    )


selected_activities = []

if beginner_activity:
    selected_activities.append("Beginner")

if intermediate_activity:
    selected_activities.append("Intermediate")

if advanced_activity:
    selected_activities.append("Advanced")


lesson_focus = st.selectbox(
    "Lesson Focus",
    [
        "Teaching New Concept",
        "Revision",
        "Practice",
        "Exam Preparation",
        "Assessment",
        "Project Lesson"
    ]
)

st.divider()


# --------------------------------------------------
# STEP 4 — TEACHING CONTEXT
# --------------------------------------------------

st.markdown("## 👨‍🏫 Step 4 — Teaching Context")

col1, col2 = st.columns(2)

with col1:
    delivery = st.selectbox(
        "Teaching Mode",
        [
            "Face-to-Face",
            "Online",
            "Hybrid"
        ]
    )

with col2:
    learning_style = st.multiselect(
        "Preferred Learning Strategy",
        [
            "Individual Learning",
            "Pair Work",
            "Group Work",
            "Collaborative Learning",
            "Think-Pair-Share",
            "Problem-Based Learning",
            "Inquiry-Based Learning"
        ]
    )

resources = st.multiselect(
    "Available Resources",
    [
        "Projector",
        "Internet",
        "Whiteboard",
        "Printed Handouts",
        "AI Tools",
        "Smart Classroom"
    ]
)

st.divider()


# --------------------------------------------------
# STEP 5 — BLOOM'S TAXONOMY
# --------------------------------------------------

st.markdown("## 🎯 Step 5 — Bloom's Taxonomy")

st.caption(
    "Select the cognitive levels you want the lesson to target."
)

blooms = st.multiselect(
    "Bloom's Levels",
    [
        "Remember",
        "Understand",
        "Apply",
        "Analyze",
        "Evaluate",
        "Create"
    ],
    default=[
        "Remember",
        "Understand",
        "Apply",
        "Analyze",
        "Evaluate",
        "Create"
    ]
)

st.divider()


# --------------------------------------------------
# STEP 6 — AI PROMPT
# --------------------------------------------------

st.markdown("## 🤖 Step 6 — AI Teaching Instructions")

teacher_prompt = st.text_area(
    "Additional Instructions (Optional)",
    height=220,
    placeholder="""
Example

Develop an interactive lesson.

Include collaborative learning.

Use authentic Computer Science examples.

Include formative assessment.

Add an exit ticket.

Provide three differentiated activities.

Align everything with Bloom's Taxonomy.
"""
)


# --------------------------------------------------
# PROMPT QUALITY
# --------------------------------------------------

word_count = len(teacher_prompt.split())

if word_count == 0:
    st.error("Prompt Quality: Not Provided")

elif word_count < 10:
    st.warning("🟥 Prompt Quality: Basic")

elif word_count < 35:
    st.info("🟨 Prompt Quality: Good")

else:
    st.success("🟩 Prompt Quality: Excellent")

st.divider()


# --------------------------------------------------
# GENERATE LESSON PLAN
# --------------------------------------------------

generate = st.button(
    "🚀 Generate Professional Lesson Plan",
    use_container_width=True
)

if generate:
    if not selected_activities:
        st.error(
            "Please select at least one activity level."
        )

    elif not blooms:
        st.error(
            "Please select at least one Bloom's Taxonomy level."
        )

    else:
        with st.spinner(
            "Generating your professional lesson plan..."
        ):
            try:
                prompt = build_prompt(
                    course,
                    programme,
                    year,
                    duration,
                    class_size,
                    proficiency,
                    course_objectives,
                    clos,
                    skill,
                    topic,
                    lesson_focus,
                    delivery,
                    learning_style,
                    resources,
                    blooms,
                    teacher_prompt,
                    selected_activities,
                    activity_bullet_points,
                    detailed_table
                )

                lesson = generate_lesson(prompt)

                st.session_state["generated_lesson"] = lesson

                st.success(
                    "Lesson Generated Successfully!"
                )

            except Exception as error:
                st.error(
                    "The lesson plan could not be generated. "
                    f"Error: {error}"
                )


# --------------------------------------------------
# DISPLAY AND DOWNLOAD GENERATED LESSON
# --------------------------------------------------

if "generated_lesson" in st.session_state:
    lesson = st.session_state["generated_lesson"]

    st.markdown("---")

    st.markdown("## 📚 Generated Lesson Plan")

    st.markdown(lesson)

    try:
        word_file = markdown_to_word(lesson)
        pdf_file = markdown_to_pdf(lesson)

        download_col1, download_col2 = st.columns(2)

        with download_col1:
            st.download_button(
                label="📘 Download Lesson Plan as Word",
                data=word_file,
                file_name="ELT_Lesson_Plan.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                use_container_width=True
            )

        with download_col2:
            st.download_button(
                label="📕 Download Lesson Plan as PDF",
                data=pdf_file,
                file_name="ELT_Lesson_Plan.pdf",
                mime="application/pdf",
                use_container_width=True
            )

    except Exception as export_error:
        st.warning(
            "The lesson plan was generated, but the Word or PDF "
            f"file could not be created. Error: {export_error}"
        )
